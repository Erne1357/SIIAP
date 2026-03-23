# app/services/document_generation_service.py
"""
Servicio de generación de documentos a partir de plantillas configurables.

Flujo:
  1. El admin sube una plantilla HTML o DOCX con marcadores {{variable}}.
  2. Al generar, el sistema sustituye los marcadores con datos del estudiante.
  3. Para HTML → PDF via weasyprint (requiere libcairo/libpango en el sistema).
  4. Para DOCX → descarga directa del archivo rellenado.

Variables disponibles:
  {{student_name}}, {{student_curp}}, {{student_email}},
  {{program_name}}, {{program_level}},
  {{period_code}}, {{period_name}},
  {{acceptance_date}}, {{coordinator_name}}, {{current_date}},
  {{control_number}}
"""
import locale
import os
from typing import Optional

from app.utils.datetime_utils import now_local


# Nombre de los meses en español (fallback si locale no está disponible)
_MESES = [
    '', 'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
    'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre',
]


def _fecha_larga(dt) -> str:
    """Retorna fecha en formato '15 de enero de 2026'."""
    return f"{dt.day} de {_MESES[dt.month]} de {dt.year}"


class DocumentGenerationService:

    # Mapa de variables → funciones (user, program, academic_period) → valor
    VARIABLE_MAP = {
        'student_name':     lambda u, p, ap, up: (
            f"{u.first_name} {u.last_name} {u.mother_last_name or ''}".strip()
        ),
        'student_curp':     lambda u, p, ap, up: getattr(u, 'curp', '') or '',
        'student_email':    lambda u, p, ap, up: u.email,
        'program_name':     lambda u, p, ap, up: p.name,
        'program_level':    lambda u, p, ap, up: p.program_level or 'Posgrado',
        'period_code':      lambda u, p, ap, up: ap.code if ap else '',
        'period_name':      lambda u, p, ap, up: ap.name if ap else '',
        'acceptance_date':  lambda u, p, ap, up: _fecha_larga(now_local()),
        'coordinator_name': lambda u, p, ap, up: (
            f"{p.coordinator.first_name} {p.coordinator.last_name}"
            if p.coordinator else ''
        ),
        'current_date':     lambda u, p, ap, up: now_local().strftime('%d/%m/%Y'),
        'control_number':   lambda u, p, ap, up: getattr(u, 'control_number', '') or '',
    }

    @staticmethod
    def get_variables(user, program, academic_period=None, user_program=None) -> dict:
        """Genera el diccionario de variables para una plantilla."""
        return {
            name: fn(user, program, academic_period, user_program)
            for name, fn in DocumentGenerationService.VARIABLE_MAP.items()
        }

    @staticmethod
    def generate(
        user_id: int,
        program_id: int,
        document_type: str,
        period_id: Optional[int] = None,
    ) -> dict:
        """
        Genera un documento para el estudiante indicado.

        Returns:
            {
                'output_path': str,   # ruta absoluta al archivo generado
                'file_type':   str,   # 'pdf' o 'docx'
                'filename':    str,   # nombre sugerido para descarga
            }

        Raises:
            ValueError: si no hay plantilla configurada o el archivo no existe.
            RuntimeError: si la generación falla.
        """
        from app.models.user import User
        from app.models.program import Program
        from app.models.academic_period import AcademicPeriod
        from app.models.user_program import UserProgram
        from app.models.document_template import DocumentTemplate
        from flask import current_app

        user = User.query.get_or_404(user_id)
        program = Program.query.get_or_404(program_id)
        period = AcademicPeriod.query.get(period_id) if period_id else None
        user_program = UserProgram.query.filter_by(
            user_id=user_id, program_id=program_id
        ).first()

        template = DocumentTemplate.get_for_program(program_id, document_type)
        if not template:
            raise ValueError(
                f"No hay plantilla activa de tipo '{document_type}' "
                f"para el programa {program_id} ni una global."
            )

        base_dir = current_app.config.get(
            'TEMPLATES_SYS_FOLDER',
            os.path.join(current_app.instance_path, 'templates_sys'),
        )
        template_abs = os.path.join(base_dir, template.file_path)
        if not os.path.exists(template_abs):
            raise ValueError(
                f"El archivo de plantilla no existe en disco: {template.file_path}"
            )

        variables = DocumentGenerationService.get_variables(
            user, program, period, user_program
        )

        output_dir = os.path.join(
            current_app.instance_path, 'uploads', 'generated', str(user_id)
        )
        os.makedirs(output_dir, exist_ok=True)

        timestamp = now_local().strftime('%Y%m%d_%H%M%S')
        safe_type = document_type.replace('_', '-')

        if template.file_type == 'html':
            output_path, filename = DocumentGenerationService._from_html(
                template_abs, variables, output_dir, safe_type, timestamp
            )
            return {'output_path': output_path, 'file_type': 'pdf', 'filename': filename}

        if template.file_type == 'docx':
            output_path, filename = DocumentGenerationService._from_docx(
                template_abs, variables, output_dir, safe_type, timestamp
            )
            return {'output_path': output_path, 'file_type': 'docx', 'filename': filename}

        raise ValueError(f"Tipo de plantilla no soportado: {template.file_type}")

    # ──────────────────────────────────────────────────────────────────────────
    # Métodos privados
    # ──────────────────────────────────────────────────────────────────────────

    @staticmethod
    def _from_html(template_abs, variables, output_dir, doc_type, timestamp):
        """Renderiza plantilla HTML con variables y convierte a PDF via weasyprint."""
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            raise RuntimeError(
                "weasyprint no está instalado. "
                "Ejecuta: pip install weasyprint"
            )

        with open(template_abs, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # Sustituir marcadores {{variable}}
        for var_name, var_value in variables.items():
            html_content = html_content.replace(f'{{{{{var_name}}}}}', str(var_value))

        filename = f"{doc_type}_{timestamp}.pdf"
        output_path = os.path.join(output_dir, filename)

        HTML(string=html_content, base_url=os.path.dirname(template_abs)).write_pdf(
            output_path
        )
        return output_path, filename

    @staticmethod
    def _from_docx(template_abs, variables, output_dir, doc_type, timestamp):
        """Rellena plantilla DOCX sustituyendo marcadores {{variable}} en párrafos y tablas."""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError(
                "python-docx no está instalado. "
                "Ejecuta: pip install python-docx"
            )

        doc = Document(template_abs)

        def replace_in_text(text):
            for var_name, var_value in variables.items():
                text = text.replace(f'{{{{{var_name}}}}}', str(var_value))
            return text

        # Párrafos del documento principal
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = replace_in_text(run.text)

        # Celdas de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = replace_in_text(run.text)

        # Encabezados y pies de página
        for section in doc.sections:
            for header_para in section.header.paragraphs:
                for run in header_para.runs:
                    run.text = replace_in_text(run.text)
            for footer_para in section.footer.paragraphs:
                for run in footer_para.runs:
                    run.text = replace_in_text(run.text)

        filename = f"{doc_type}_{timestamp}.docx"
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)
        return output_path, filename
